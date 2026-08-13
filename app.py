import os
import re
import io
import uuid
import json
import base64 as _b64
import tempfile
import random
import hashlib
import hmac
import time
import smtplib
from email.mime.text import MIMEText
from email.header import Header
from email.utils import formataddr

from functools import wraps

import requests
from flask import Flask, request, send_file, jsonify, render_template
import genanki

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

DEEPSEEK_URL = "https://api.deepseek.com/v1/chat/completions"


def load_env_file():
    """手写 .env 加载：DEEPSEEK_API_KEY=xxx（不额外装 python-dotenv）"""
    p = os.path.join(BASE_DIR, ".env")
    if not os.path.exists(p):
        return
    try:
        with open(p, encoding="utf-8", errors="ignore") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                k, v = line.split("=", 1)
                os.environ.setdefault(k.strip(), v.strip().strip('"').strip("'"))
    except Exception:
        pass


def load_deepseek_key():
    """key 优先级：环境变量 > .env 文件"""
    load_env_file()
    k = os.environ.get("DEEPSEEK_API_KEY")
    return k if k and not k.startswith("PUT_YOUR") else None


def load_zhipu_key():
    """ZHIPU_API_KEY：环境变量 > .env > Windows 用户级环境变量（注册表）兜底"""
    load_env_file()
    k = os.environ.get("ZHIPU_API_KEY")
    if k and not k.startswith("PUT_YOUR"):
        return k
    try:
        import winreg
        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, "Environment") as key:
            val, _ = winreg.QueryValueEx(key, "ZHIPU_API_KEY")
            if val and not str(val).startswith("PUT_YOUR"):
                return str(val).strip()
    except Exception:
        pass
    return None


DEEPSEEK_KEY = load_deepseek_key()
ZHIPU_KEY = load_zhipu_key()
ZHIPU_URL = "https://open.bigmodel.cn/api/paas/v4/chat/completions"

SYSTEM_PROMPT = """你是一名医学教育内容专家，擅长将教材/讲义内容转化为适合间隔重复(Anki)的原子化记忆卡片。
要求：
1. 每张卡只考查一个独立知识点（原子化），避免复合问题。
2. front 为问题或概念提示；back 用三段式：「答案：」精炼要点，「解析：」一句话解释为什么，「关联考点：」标注所属科目章节（如 生理·肾脏滤过）。段落间用换行分隔。
3. 覆盖关键维度：定义、机制、鉴别诊断、重要数值、临床意义。
4. 医学准确性优先；对不确定的内容标注「需核实」。
5. 适度添加 tags，如 ["306","生理","循环"]。
6. 按内容密度生成 5-20 张卡片，宁缺毋滥。
7. 每张卡附带一个「选择题版」quiz 字段：{"question":"基于本卡知识点的题干","options":["选项A","选项B","选项C","选项D","选项E"],"answer":正确选项索引(0-4)}。仿考研真题五选一（A-E），选项简短清晰，干扰项合理有迷惑性。
8. 每张卡必须带 source 字段：**原样引用**输入资料中最能支撑该卡答案的一句话（不超过 80 字，不得改写或概括），用于用户对照原文核实。
9. 所有字符串值内部严禁使用英文双引号（会破坏 JSON），引用术语用「」。
只输出 JSON，结构：{"cards":[{"front":"...","back":"...","tags":["..."],"quiz":{"question":"...","options":["..."],"answer":0},"source":"..."}]}"""

WRONG_PROMPT = """你是一名医学教育内容专家，专门帮助考研学生「把做错的题变成记忆卡」，下次不再错。
用户会粘贴一道 TA 做错的题（题干、选项、TA 的答案或当时纠结的点）。
要求：
1. 从错题中提炼背后的知识点，生成记忆卡——front 是「针对该考点的问题」（不是复述原题，而是考查题目背后的知识点）。
2. back 用三段式：「答案：」该考点的正确结论，「解析：」点破常见错误认知/为什么容易选错（结合用户错因），「关联考点：」所属科目章节（如 内科学·呼吸·COPD）。段落间用换行分隔。
3. 若题目涉及易混概念或鉴别诊断，优先做成「对比卡」（如 A vs B 的关键区别）。
4. tags 必须含「错题」，再加科目/章节标签，如 ["错题","306","内科学"]。
5. 医学准确性优先；对不确定的内容标注「需核实」。
6. 生成 1-5 张，宁缺毋滥——一道错题通常 1-2 张就够。
7. 每张卡附带一个「选择题版」quiz 字段：{"question":"基于该考点的题干","options":["选项A","选项B","选项C","选项D","选项E"],"answer":正确选项索引(0-4)}。仿考研真题五选一（A-E），选项简短清晰，干扰项合理有迷惑性。
8. 每张卡必须带 source 字段：**原样引用**用户粘贴的错题中最能支撑该卡答案的一句话（不超过 80 字，不得改写）。
9. 所有字符串值内部严禁使用英文双引号（会破坏 JSON），引用术语用「」。
只输出 JSON，结构：{"cards":[{"front":"...","back":"...","tags":["..."],"quiz":{"question":"...","options":["..."],"answer":0},"source":"..."}]}"""

MODEL_ID = 1607392319

# Anki 卡片模型：结构固定，模块级建一次所有请求复用
ANKI_MODEL = genanki.Model(
    MODEL_ID,
    "MedCardModel",
    fields=[{"name": "Front"}, {"name": "Back"}],
    templates=[
        {
            "name": "Card",
            "qfmt": "{{Front}}",
            "afmt": '{{Front}}<hr id="answer">{{Back}}',
        }
    ],
)

# ---------- API 鉴权：页面加载时按 IP+小时签发短 token（HMAC），有效期至下一小时结束 ----------
# 设计：token 由 GET / 渲染页面时注入 HTML（data-token），前端 fetch 带 X-API-Token 头。
# 陌生人不知道 API_SECRET 就算拿到接口地址也算不出 token → 401。
# 未配置 API_SECRET（.env/环境变量均无）时鉴权自动关闭（本地开发不锁死）。
# ponytail: 无登录系统的 MVP 能防的最懒方案；不做每请求刷新/限流，接口被反代转发时以 remote_addr 为准。
API_SECRET = os.environ.get("API_SECRET", "")


def make_token(ip, hour):
    """HMAC-SHA256(secret, 'ip:小时') 前 16 位。小时粒度 = 同一小时多次刷新页面拿到同一 token。"""
    msg = f"{ip}:{hour}".encode()
    return hmac.new((API_SECRET or "nosecret").encode(), msg, hashlib.sha256).hexdigest()[:16]


def token_ok():
    """校验 X-API-Token。当前小时与上一小时都接受（跨小时边界容忍，避免整点报错）。"""
    t = request.headers.get("X-API-Token", "")
    if not t:
        return False
    ip = request.remote_addr or "0.0.0.0"
    hour = int(time.time()) // 3600
    return any(hmac.compare_digest(t, make_token(ip, h)) for h in (hour, hour - 1))


def api_guard(fn):
    """AI/邮件接口守卫：无有效 token 一律 401（在业务校验之前执行）"""

    @wraps(fn)
    def wrapper(*a, **k):
        if API_SECRET and not token_ok():
            return jsonify({"error": "未授权：请从网站页面操作"}), 401
        return fn(*a, **k)

    return wrapper


# ---------- 意见反馈邮件（163 SMTP，授权码走环境变量 SMTP_AUTH_CODE） ----------
SMTP_HOST = "smtp.163.com"
SMTP_PORT = 465
SMTP_USER = "suqiqinak@163.com"
SMTP_AUTH_CODE = os.environ.get("SMTP_AUTH_CODE", "")


def send_feedback_email(msg):
    """发意见反馈邮件到作者邮箱；msg 为纯文本。无授权码或发送失败抛异常。"""
    if not SMTP_AUTH_CODE:
        raise RuntimeError("未配置 SMTP_AUTH_CODE（.env 或环境变量）")
    mime = MIMEText(msg, "plain", "utf-8")
    mime["From"] = formataddr((str(Header("医学记忆卡反馈", "utf-8")), SMTP_USER))
    mime["To"] = formataddr((str(Header("作者", "utf-8")), SMTP_USER))
    mime["Subject"] = str(Header("【医学记忆卡】用户意见反馈", "utf-8"))
    with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, timeout=15) as srv:
        srv.login(SMTP_USER, SMTP_AUTH_CODE)
        srv.sendmail(SMTP_USER, [SMTP_USER], mime.as_string())


def parse_ai_json(content):
    """容错解析 AI 返回的 JSON：先直解，失败截取最外层花括号再解，仍失败返回 None。"""
    if not content:
        return None
    try:
        return json.loads(content)
    except Exception:
        pass
    i, j = content.find("{"), content.rfind("}")
    if i < 0 or j <= i:
        return None
    s = content[i: j + 1]
    try:
        return json.loads(s)
    except Exception:
        return None


def ai_extract_cards(text, max_cards=10, mode="text", custom_tags=None):
    """调用 DeepSeek 抽卡；无 key 时返回 None（走演示模式）。mode='wrong' 为错题制卡。"""
    if not DEEPSEEK_KEY:
        return None
    if mode == "wrong":
        base_prompt = WRONG_PROMPT.replace(
            "6. 生成 1-5 张，宁缺毋滥——一道错题通常 1-2 张就够。",
            f"6. 生成最多 {max_cards} 张，宁缺毋滥——一道错题通常 1-2 张就够。",
        )
        user_msg = f"这是我做错的一道题，请帮我提炼知识点记忆卡：\n\n{text}"
    else:
        base_prompt = SYSTEM_PROMPT.replace(
            "按内容密度生成 5-20 张卡片，宁缺毋滥。",
            f"按内容密度生成最多 {max_cards} 张卡片，宁缺毋滥。",
        )
        user_msg = f"请将以下医学教材/讲义内容转化为记忆卡片：\n\n{text}"
    # 用户自定义标签注入 Prompt：AI 生成 tags 时优先匹配用户标签体系
    if custom_tags:
        tags_str = "、".join(str(t) for t in custom_tags if str(t).strip())
        if tags_str:
            base_prompt += (
                f"\n补充要求：用户自定义了标签体系：【{tags_str}】。"
                "生成 tags 时，若卡片知识点与其中某个标签匹配，必须使用该标签（可与其他标签并存）；"
                "确实不匹配时照常按科目/章节打标签。"
            )
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_KEY}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": base_prompt},
            {"role": "user", "content": user_msg},
        ],
        "response_format": {"type": "json_object"},
        "temperature": 0.3,
    }
    resp = requests.post(DEEPSEEK_URL, headers=headers, json=payload, timeout=120)
    resp.raise_for_status()
    content = resp.json()["choices"][0]["message"]["content"]
    data = parse_ai_json(content)
    if not data:
        # 偶发坏 JSON：原样重试一次
        resp = requests.post(DEEPSEEK_URL, headers=headers, json=payload, timeout=120)
        resp.raise_for_status()
        content = resp.json()["choices"][0]["message"]["content"]
        data = parse_ai_json(content)
    return (data or {}).get("cards", [])


def rule_extract_cards(text, max_cards=15):
    """演示模式：句子级切分生成卡（无 key 时兜底，质量有限）。"""
    sentences = re.split(r"(?<=[。；;.!?！？])\s*", text)
    cards = []
    for s in sentences:
        s = s.strip()
        if len(s) < 10:
            continue
        hint = s[:12] + ("…" if len(s) > 12 else "")
        cards.append(
            {
                "front": f"请回忆/解释：{hint}",
                "back": s,
                "tags": ["演示", "规则抽卡"],
                "source": s[:80],
            }
        )
        if len(cards) >= max_cards:
            break
    return cards


def build_deck(cards, deck_name):
    deck_id = random.randint(1 << 30, 1 << 31)
    deck = genanki.Deck(deck_id, deck_name or "医学记忆卡")
    for c in cards:
        note = genanki.Note(
            ANKI_MODEL, fields=[str(c.get("front", "")), str(c.get("back", ""))]
        )
        if c.get("tags"):
            note.tags = [str(t) for t in c["tags"]]
        deck.add_note(note)
    fname = f"{(deck_name or 'deck')}_{uuid.uuid4().hex[:8]}.apkg"
    # 内存生成（临时文件同请求内使用，兼容 serverless 无持久文件系统）
    fd, tmp = tempfile.mkstemp(suffix=".apkg")
    os.close(fd)
    try:
        genanki.Package(deck).write_to_file(tmp)
        with open(tmp, "rb") as fh:
            data = fh.read()
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)
    return data, fname


app = Flask(__name__)


@app.route("/")
def index():
    # 页面注入当小时 token，前端 fetch 携带 X-API-Token 头
    ip = request.remote_addr or "0.0.0.0"
    return render_template("index.html", api_token=make_token(ip, int(time.time()) // 3600))


@app.route("/records")
def records():
    return render_template("records.html")


@app.route("/terms")
def terms():
    # 用户服务协议：生成式 AI 合规（标识办法§8/暂行办法§9-11,15）
    return render_template("terms.html")


@app.route("/generate", methods=["POST"])
@api_guard
def generate():
    print(f">>> /generate key={'SET' if DEEPSEEK_KEY else 'NONE'}", flush=True)
    data = request.get_json(force=True, silent=True) or {}
    text = (data.get("text") or "").strip()
    if len(text) > 30000:
        text = text[:30000]  # 超长截断：防 DeepSeek 超时/超上下文（前端同步限制）
    deck_name = (data.get("deck_name") or "医学记忆卡").strip()[:40]
    if not text:
        return jsonify({"error": "请输入或上传内容"}), 400
    try:
        try:
            max_cards = int(data.get("max_cards") or 10)
        except (TypeError, ValueError):
            max_cards = 10
        max_cards = max(1, min(20, max_cards))
        mode = "wrong" if data.get("mode") == "wrong" else "text"
        custom_tags = data.get("custom_tags") or []
        if not isinstance(custom_tags, list):
            custom_tags = []
        custom_tags = [str(t)[:20] for t in custom_tags][:30]
        cards = ai_extract_cards(text, max_cards, mode=mode, custom_tags=custom_tags)
        demo = False
        if cards is None:
            demo = True
            cards = rule_extract_cards(text, max_cards)
            if mode == "wrong":
                for c in cards:
                    if "错题" not in c.get("tags", []):
                        c["tags"] = ["错题"] + c.get("tags", [])
        if not cards:
            return jsonify({"error": "未能生成卡片，请检查内容或 API Key"}), 400
        data, fname = build_deck(cards, deck_name)
        return jsonify(
            {
                "file": fname,
                "file_b64": _b64.b64encode(data).decode(),
                "cards": len(cards),
                "demo": demo,
                "cards_content": cards,
            }
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False


MAX_UPLOAD = 4 * 1024 * 1024  # 4MB（Vercel 函数体 4.5MB 限制的保守值）


@app.route("/extract_pdf", methods=["POST"])
@api_guard
def extract_pdf():
    """上传 PDF → 提取文字；扫描件（无文字层）自动渲染成图片走智谱 OCR 兜底（最多前 3 页）"""
    f = request.files.get("file")
    if not f or not f.filename.lower().endswith(".pdf"):
        return jsonify({"error": "请上传 PDF 文件"}), 400
    try:
        data = f.read()
        if len(data) > MAX_UPLOAD:
            return jsonify({"error": "文件过大（限 4MB）"}), 400
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            t = page.extract_text() or ""
            if t.strip():
                parts.append(t)
        text = "\n".join(parts).strip()
        if not text:
            # 扫描件/图片版 PDF → 渲染成图片走 AI 识别
            if not ZHIPU_KEY:
                return jsonify({"error": "未能提取到文字（扫描件 PDF），且未配置 ZHIPU_API_KEY 无法 AI 识别"}), 400
            try:
                import pymupdf
            except ImportError:
                return jsonify({"error": "未能提取到文字（扫描件 PDF），且服务器缺少渲染库"}), 500
            doc = pymupdf.open(stream=data, filetype="pdf")
            ocr_parts = []
            for i, page in enumerate(doc):
                if i >= 3:
                    break
                pix = page.get_pixmap(dpi=150)
                png = pix.tobytes("png")
                ocr_parts.append(zhipu_ocr_image(png, "image/png"))
            text = "\n".join(ocr_parts).strip()
            if not text:
                return jsonify({"error": "扫描件 AI 识别也未提取到文字"}), 400
            return jsonify({"text": text[:50000], "ocr": True})
        return jsonify({"text": text[:50000]})
    except Exception as e:
        return jsonify({"error": f"PDF 解析失败：{e}"}), 500


@app.route("/extract_docx", methods=["POST"])
@api_guard
def extract_docx():
    """上传 Word(.docx) → 提取段落+表格文字 → 返回（限 5 万字）"""
    f = request.files.get("file")
    if not f or not f.filename.lower().endswith(".docx"):
        return jsonify({"error": "请上传 .docx 文件"}), 400
    try:
        from docx import Document
        data = f.read()
        if len(data) > MAX_UPLOAD:
            return jsonify({"error": "文件过大（限 4MB）"}), 400
        doc = Document(io.BytesIO(data))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
        if not text:
            return jsonify({"error": "未能提取到文字（可能是图片型 Word）"}), 400
        return jsonify({"text": text[:50000]})
    except Exception as e:
        return jsonify({"error": f"Word 解析失败：{e}"}), 500


IMG_EXTS = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".webp": "image/webp"}


def clean_ocr_text(text):
    """OCR 后处理：规则清洗 App 界面噪音行（确定性兜底，不依赖模型心情）"""
    drop_prefix = ("统计", "标签", "来源", "难度", "本题", "全部考生", "本人答", "正确率")
    drop_kw = ("纠错", "1.1万", "写评论", "收藏", "点赞", "评论", "笔记", "有争议", "已过时")
    lines = []
    for ln in (text or "").split("\n"):
        s = ln.strip()
        if not s:
            continue
        if s.startswith(drop_prefix):
            continue
        if any(k in s for k in drop_kw):
            continue
        lines.append(s)
    return "\n".join(lines)


def zhipu_ocr_image(data, mime):
    """智谱免费视觉模型识别图片文字 → 返回清洗后的文本（无 key/失败抛异常）"""
    if not ZHIPU_KEY:
        raise RuntimeError("未配置 ZHIPU_API_KEY")
    b64 = _b64.b64encode(data).decode()
    payload = {
        "model": "glm-4v-flash",
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                    {"type": "text", "text": "请提取这张图片中的医学题目内容。\n如果图片是练习/考试 App 的题目截图：提取「题干」「选项（A/B/C/D/E 及内容）」「答案与解析（如有）」，可保留顶部科目/章节信息（如：外科学 第二十三章 乳房疾病）和题型标注（如 A3/A4 型题）。\n必须忽略所有与题目无关的界面元素：状态栏（时间/信号/电量）、导航栏、搜索栏、按钮文字（写评论/笔记/收藏/评论/点赞）、「5条纠错」「难度」「标签」「来源」「1.1万」等。特别注意：任何以「统计：」开头的整行（包含收藏数/作答次数/正确率/本人答等数字）必须整行丢弃。\n如果图片不是题目截图，则提取全部文字。\n原样输出，保持段落与选项分行，不要添加任何解释、评论或格式标记。如果图片里没有文字，只输出：无文字"},
                ],
            }
        ],
        "temperature": 0.1,
    }
    resp = None
    for attempt in range(2):
        try:
            resp = requests.post(ZHIPU_URL, headers={"Authorization": f"Bearer {ZHIPU_KEY}"}, json=payload, timeout=120)
            resp.raise_for_status()
            break
        except Exception:
            # 免费模型偶发限流/抖动：重试一次，仍失败则抛给上层
            if attempt == 1:
                raise
    content = resp.json()["choices"][0]["message"]["content"].strip()
    return clean_ocr_text(content)


@app.route("/ocr", methods=["POST"])
@api_guard
def ocr():
    """上传图片 → 智谱免费视觉模型(glm-4v-flash)提取文字 → 返回（限 4MB）"""
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "请上传图片"}), 400
    ext = os.path.splitext(f.filename or "")[1].lower()
    mime = IMG_EXTS.get(ext)
    if not mime:
        return jsonify({"error": "仅支持 png / jpg / jpeg / webp 图片"}), 400
    if not ZHIPU_KEY:
        return jsonify({"error": "未配置 ZHIPU_API_KEY，无法识别图片（可在 .env 填入或配置系统环境变量）"}), 400
    data = f.read()
    if len(data) > MAX_UPLOAD:
        return jsonify({"error": "图片过大（限 4MB）"}), 400
    try:
        return jsonify({"text": zhipu_ocr_image(data, mime)[:50000]})
    except Exception as e:
        return jsonify({"error": f"图片识别失败：{e}"}), 500


@app.route("/feedback", methods=["POST"])
@api_guard
def feedback():
    """意见反馈：前端提交文本 → 163 SMTP 发到作者邮箱（无授权码时返回 501 提示）"""
    data = request.get_json(force=True, silent=True) or {}
    msg = (data.get("msg") or "").strip()
    if not msg:
        return jsonify({"error": "反馈内容不能为空"}), 400
    if len(msg) > 2000:
        msg = msg[:2000]
    try:
        send_feedback_email("用户意见反馈：\n\n" + msg)
        return jsonify({"ok": True})
    except RuntimeError as e:
        return jsonify({"error": str(e)}), 501
    except Exception as e:
        return jsonify({"error": f"邮件发送失败：{e}"}), 500


if __name__ == "__main__":
    print(
        f">>> STARTUP key={'SET' if DEEPSEEK_KEY else 'NONE'} "
        f"zhipu={'SET' if ZHIPU_KEY else 'NONE'} "
        f"smtp={'SET' if SMTP_AUTH_CODE else 'NONE'}",
        flush=True,
    )
    app.run(host="127.0.0.1", port=5000, debug=False)
